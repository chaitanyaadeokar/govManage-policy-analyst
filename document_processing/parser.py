"""
Document Parser - Extract text and structure from uploaded policy/framework documents.
Supports PDF, DOCX, TXT, and more.
"""
import os
import io
from typing import Dict, Any, List, Optional
from pathlib import Path
import mimetypes

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠ PyPDF2 not installed. Install with: pip install PyPDF2")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠ python-docx not installed. Install with: pip install python-docx")


class DocumentParser:
    """
    Parse various document formats to extract text content.
    Supports: PDF, DOCX, TXT, MD
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md', '.doc']
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document file and extract text content.
        
        Args:
            file_path: Path to the document file
        
        Returns:
            Dict with extracted text, metadata, and structure
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {self.supported_formats}")
        
        # Parse based on file type
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            return self._parse_text(file_path)
        else:
            raise ValueError(f"Parser not implemented for: {file_ext}")
    
    def parse_bytes(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse document from bytes (for API uploads).
        
        Args:
            file_bytes: Document content as bytes
            filename: Original filename (to determine type)
        
        Returns:
            Dict with extracted text and metadata
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return self._parse_pdf_bytes(file_bytes)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx_bytes(file_bytes)
        elif file_ext in ['.txt', '.md']:
            return self._parse_text_bytes(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        with open(file_path, 'rb') as f:
            return self._parse_pdf_bytes(f.read(), str(file_path))
    
    def _parse_pdf_bytes(self, file_bytes: bytes, filename: str = "document.pdf") -> Dict[str, Any]:
        """Parse PDF from bytes."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed")
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        
        # Extract metadata
        metadata = {
            "num_pages": len(pdf_reader.pages),
            "filename": filename,
            "format": "pdf"
        }
        
        # Try to get PDF metadata
        if pdf_reader.metadata:
            metadata.update({
                "title": pdf_reader.metadata.get('/Title', ''),
                "author": pdf_reader.metadata.get('/Author', ''),
                "subject": pdf_reader.metadata.get('/Subject', ''),
                "creator": pdf_reader.metadata.get('/Creator', '')
            })
        
        # Extract text from all pages
        pages = []
        full_text = []
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                text = page.extract_text()
                pages.append({
                    "page_number": page_num,
                    "text": text,
                    "char_count": len(text)
                })
                full_text.append(text)
            except Exception as e:
                print(f"⚠ Error extracting page {page_num}: {e}")
                pages.append({
                    "page_number": page_num,
                    "text": "",
                    "error": str(e)
                })
        
        return {
            "text": "\n\n".join(full_text),
            "metadata": metadata,
            "pages": pages,
            "total_chars": sum(len(p.get("text", "")) for p in pages),
            "extraction_method": "PyPDF2"
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        with open(file_path, 'rb') as f:
            return self._parse_docx_bytes(f.read(), str(file_path))
    
    def _parse_docx_bytes(self, file_bytes: bytes, filename: str = "document.docx") -> Dict[str, Any]:
        """Parse DOCX from bytes."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        doc = DocxDocument(io.BytesIO(file_bytes))
        
        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "filename": filename,
            "format": "docx",
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else ""
        }
        
        # Extract paragraphs
        paragraphs = []
        full_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                })
                full_text.append(text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "text": "\n\n".join(full_text),
            "metadata": metadata,
            "paragraphs": paragraphs,
            "tables": tables,
            "total_chars": sum(len(p["text"]) for p in paragraphs),
            "extraction_method": "python-docx"
        }
    
    def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return self._parse_text_bytes(text.encode('utf-8'), str(file_path))
    
    def _parse_text_bytes(self, file_bytes: bytes, filename: str = "document.txt") -> Dict[str, Any]:
        """Parse text from bytes."""
        text = file_bytes.decode('utf-8', errors='ignore')
        
        lines = text.split('\n')
        
        return {
            "text": text,
            "metadata": {
                "filename": filename,
                "format": "text",
                "line_count": len(lines)
            },
            "lines": lines,
            "total_chars": len(text),
            "extraction_method": "text"
        }
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks for processing.
        
        Args:
            text: Text to chunk
            chunk_size: Target size of each chunk (characters)
            overlap: Overlap between chunks (characters)
        
        Returns:
            List of chunks with metadata
        """
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            chunks.append({
                "chunk_id": chunk_id,
                "text": chunk_text,
                "start_pos": start,
                "end_pos": end,
                "char_count": len(chunk_text)
            })
            
            chunk_id += 1
            start = end - overlap
        
        return chunks


# Global parser instance
document_parser = DocumentParser()


if __name__ == "__main__":
    # Test the parser
    print("Document Parser Test")
    print("=" * 60)
    
    # Test with a sample text file
    test_text = """
    Sample Compliance Framework
    
    Article 1: Data Protection
    All personal data must be protected using encryption.
    
    Article 2: Access Control
    Only authorized personnel may access sensitive data.
    """
    
    result = document_parser.parse_bytes(test_text.encode(), "test.txt")
    print(f"✓ Parsed text document")
    print(f"  Characters: {result['total_chars']}")
    print(f"  Lines: {result['metadata']['line_count']}")
    
    # Test chunking
    chunks = document_parser.chunk_text(test_text, chunk_size=100, overlap=20)
    print(f"\n✓ Created {len(chunks)} chunks")
